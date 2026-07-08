from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 全局样式
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_code_block(doc, code, label=""):
    if label:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(11)
    for line in code.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F5F5F5')
        shading.set(qn('w:val'), 'clear')
        p._p.get_or_add_pPr().append(shading)


# ===== 封面 =====
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SQL 注入漏洞安全审计报告')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('用户信息管理平台 — Flask Web 应用')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x77, 0x99)

doc.add_paragraph()
doc.add_paragraph()

info = [
    ('项目名称', '用户信息管理平台'),
    ('审计日期', '2026年7月8日'),
    ('审计内容', 'SQL 注入漏洞专项审计'),
    ('审计文件', 'app.py — /register 和 /search 路由'),
    ('仓库地址', 'https://github.com/nnm78956/user-management'),
]
for k, v in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{k}：{v}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ===== 1. 什么是 SQL 注入 =====
doc.add_heading('1. 什么是 SQL 注入', level=1)

doc.add_paragraph(
    'SQL 注入（SQL Injection）是一种常见的 Web 安全漏洞。'
    '攻击者通过在输入框中输入特殊的 SQL 代码，让程序在执行数据库查询时'
    '把攻击者的输入当作 SQL 命令来执行，从而获取或篡改数据库中的数据。'
)

doc.add_paragraph(
    '举个例子：如果程序直接把用户输入的内容拼接到 SQL 语句中——'
)

add_code_block(doc, '''# 有漏洞的写法（拼接字符串）
keyword = request.args.get("keyword", "")
sql = f"SELECT * FROM users WHERE username LIKE '%{keyword}%'"
c.execute(sql)''', '❌ 危险写法：')

doc.add_paragraph()
doc.add_paragraph('当用户输入 `\' OR \'1\'=\'1` 时，实际的 SQL 语句变成了：')

add_code_block(doc, '''SELECT * FROM users WHERE username LIKE '%' OR '1'='1%' ''', '实际执行的 SQL：')

doc.add_paragraph()
doc.add_paragraph(
    '由于 `OR \'1\'=\'1\'` 永远为真，这条语句会返回数据库中所有的用户信息。'
    '攻击者还可以使用 UNION 查询来读取其他表的数据，甚至修改或删除数据。'
)

doc.add_page_break()

# ===== 2. 漏洞位置 =====
doc.add_heading('2. 漏洞位置', level=1)

doc.add_paragraph('本次审计发现项目中有两处 SQL 注入漏洞：')

doc.add_heading('2.1 搜索功能（/search 路由）', level=2)

doc.add_paragraph('文件：app.py 第 130 行（修复前）')
doc.add_paragraph('问题：搜索关键词直接拼接到 SQL 查询语句中。')

add_code_block(doc, '''@app.route("/search")
def search():
    keyword = request.args.get("keyword", "")
    # ...
    if keyword:
        sql = f"SELECT id, username, email, phone FROM users
                WHERE username LIKE '%{keyword}%'
                OR email LIKE '%{keyword}%'"    # ← 直接拼接用户输入
        c.execute(sql)''', '❌ 修复前的代码：')

doc.add_paragraph()

doc.add_heading('2.2 注册功能（/register 路由）', level=2)

doc.add_paragraph('文件：app.py 第 101 行（修复前）')
doc.add_paragraph('问题：用户名、密码、邮箱、手机号全都直接拼接到 INSERT 语句中。')

add_code_block(doc, '''@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form.get("username", "")
        password = request.form.get("password", "")
        email = request.form.get("email", "")
        phone = request.form.get("phone", "")

        sql = f"INSERT INTO users (username, password, email, phone)
                VALUES ('{username}', '{password}', '{email}', '{phone}')"
                # ↑ 四个字段全部直接拼接用户输入
        c.execute(sql)''', '❌ 修复前的代码：')

doc.add_page_break()

# ===== 3. 漏洞验证 =====
doc.add_heading('3. 漏洞验证', level=1)

doc.add_paragraph(
    '下面是在修复前实际测试的结果，证明这两处漏洞确实可以被利用：'
)

doc.add_heading('3.1 搜索注入 — UNION 查询', level=2)

add_code_block(doc, '''# 攻击者输入（搜索框）：
1' UNION SELECT 1,name,3,4 FROM sqlite_master WHERE type='table'--

# 实际执行的 SQL（拼接后）：
SELECT id, username, email, phone FROM users
WHERE username LIKE '%1'
UNION SELECT 1,name,3,4 FROM sqlite_master WHERE type='table'--%'
OR email LIKE '%1'
UNION SELECT 1,name,3,4 FROM sqlite_master WHERE type='table'--%'

# 结果：成功读取到数据库中的表名 "users"''')

doc.add_paragraph()

doc.add_heading('3.2 搜索注入 — 读取密码', level=2)

add_code_block(doc, '''# 攻击者输入（搜索框）：
1' UNION SELECT 1,username,password,email FROM users--

# 结果：页面上显示了所有用户的明文密码：
# admin123、alice2025''')

doc.add_paragraph()

doc.add_heading('3.3 构造万能密码登录', level=2)
doc.add_paragraph('虽然登录功能使用的是 USERS 字典而非数据库，但注册功能的注入同样危险。')

doc.add_paragraph()
doc.add_paragraph('修复前的验证结果：', style='List Bullet')
doc.add_paragraph('✅ UNION 注入成功，可读取任意数据库内容（高危）', style='List Bullet')
doc.add_paragraph('✅ 用户密码明文泄露（高危）', style='List Bullet')
doc.add_paragraph('✅ INSERT 注入成功，可执行任意 SQL 语句（高危）', style='List Bullet')

doc.add_page_break()

# ===== 4. 修复方案 =====
doc.add_heading('4. 修复方案', level=1)

doc.add_paragraph(
    '修复 SQL 注入的方法很简单：不要用字符串拼接 SQL，而是使用参数化查询。'
    '参数化查询会把用户输入当作"数据"而不是"SQL 代码"来处理。'
)

doc.add_paragraph()

doc.add_heading('4.1 什么是参数化查询', level=2)

doc.add_paragraph(
    '参数化查询使用问号 ? 作为占位符，然后把用户输入的值通过第二个参数传进去。'
    '数据库会严格区分"SQL 语句结构"和"传入的数据"，'
    '所以即使用户输入了 SQL 关键字，也不会被当作命令执行。'
)

doc.add_paragraph()

doc.add_heading('4.2 搜索功能修复', level=2)

add_code_block(doc, '''@app.route("/search")
def search():
    keyword = request.args.get("keyword", "")
    # ...
    if keyword:
        sql = "SELECT id, username, email, phone FROM users
                WHERE username LIKE ? OR email LIKE ?"
        #    ↑ 用 ? 代替了 %{keyword}%

        like_pattern = f"%{keyword}%"
        c.execute(sql, (like_pattern, like_pattern))
        #    ↑ 数据通过第二个参数传入，不会被当作 SQL 执行''', '✅ 修复后的代码：')

doc.add_paragraph()
doc.add_paragraph()
doc.add_heading('4.3 注册功能修复', level=2)

add_code_block(doc, '''@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        # ... 获取表单数据 ...

        sql = "INSERT INTO users (username, password, email, phone)
                VALUES (?, ?, ?, ?)"
        #    ↑ 用 ? 代替了拼接

        params = (username, password, email, phone)
        c.execute(sql, params)
        #    ↑ 数据通过元组传入，安全''', '✅ 修复后的代码：')

doc.add_page_break()

# ===== 5. 修复前后对比 =====
doc.add_heading('5. 修复前后对比', level=1)

table = doc.add_table(rows=3, cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['功能', '修复前（有注入）', '修复后（安全）', '效果']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '1A1A2E')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(10)

data = [
    ['用户搜索', "f'...LIKE %{keyword}%'", '? 占位符传参',
     "' OR '1'='1 被当作普通文字搜索，不泄露数据"],
    ['用户注册', "f'...VALUES ('{username}')'", '? 占位符传参',
     '恶意 SQL 被当作用户名存入，不执行'],
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val
        for p in table.rows[i + 1].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# ===== 6. 修复验证 =====
doc.add_heading('6. 修复验证', level=1)

doc.add_paragraph('修复后，我们用同样的攻击手法再次测试：')

doc.add_paragraph()

table2 = doc.add_table(rows=5, cols=2, style='Table Grid')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

test_data = [
    ('测试内容', '测试结果'),
    ('搜索 `\' OR \'1\'=\'1`', '✅ 无搜索结果（注入被拦截）'),
    ('搜索 UNION 查询', '✅ UNION 语句被当作普通文字，不执行'),
    ('搜索读取密码', '✅ 密码字段无法被读取'),
    ('注册恶意 SQL', '✅ 恶意内容只作为用户名存储，不执行 SQL'),
]
for i, (k, v) in enumerate(test_data):
    table2.rows[i].cells[0].text = k
    table2.rows[i].cells[1].text = v
    if i == 0:
        for cell in table2.rows[0].cells:
            set_cell_shading(cell, '1A1A2E')
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True

doc.add_paragraph()

# ===== 7. 总结 =====
doc.add_heading('7. 总结', level=1)

doc.add_paragraph(
    'SQL 注入漏洞的修复其实很简单：'
)

doc.add_paragraph('永远不要用字符串拼接的方式构造 SQL 语句', style='List Bullet')
doc.add_paragraph('始终使用参数化查询（? 占位符）', style='List Bullet')
doc.add_paragraph('让数据库替你区分"代码"和"数据"', style='List Bullet')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('核心原则：')
run.bold = True
run.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)
run.font.size = Pt(12)

doc.add_paragraph(
    '用户的输入永远只是"数据"，不应该成为"SQL 代码"。'
    '参数化查询就是实现这一原则的最简单有效的方式。'
)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 报告结束 —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# 保存
path = '/home/user/user-management/sql-injection-audit-report.docx'
doc.save(path)
print(f'✅ 报告已生成: {path}')

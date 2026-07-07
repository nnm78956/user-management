from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ===========================
# 全局样式设置
# ===========================
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_colored_heading(doc, text, color_rgb, level=1):
    """添加带颜色的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color_rgb
    return heading


def add_risk_badge(doc, text):
    """添加风险等级标签"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if '高危' in text:
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    elif '中危' in text:
        run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
    elif '低危' in text:
        run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
    run.font.size = Pt(11)
    return p


def add_code_snippet(doc, code, label="原版代码"):
    """添加代码块，用灰色底色和等宽字体"""
    p = doc.add_paragraph()
    run = p.add_run(f"📄 {label}：")
    run.bold = True
    run.font.size = Pt(11)

    for line in code.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # 灰色背景通过 shading 实现
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F0F0F0')
        shading.set(qn('w:val'), 'clear')
        p._p.get_or_add_pPr().append(shading)


# ===========================
# 封面
# ===========================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('安全漏洞审计报告')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('用户信息管理平台 — Flask Web 应用')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x77, 0x99)

doc.add_paragraph()

info_items = [
    ('审计日期', '2026年7月7日'),
    ('审计类型', '静态源代码审计'),
    ('审计工具', 'Claude Code (Anthropic)'),
    ('审计版本', 'v1（原版）→ v2（修复版）'),
    ('仓库地址（原版）', 'github.com/nnm78956/user-management-original'),
    ('仓库地址（修复版）', 'github.com/nnm78956/user-management'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ===========================
# 目录页
# ===========================
add_colored_heading(doc, '目  录', RGBColor(0x1A, 0x1A, 0x2E), level=1)
toc_items = [
    '1. 审计概述',
    '2. 漏洞详情',
    '    2.1 漏洞一：密码明文存储 🔴 高危',
    '    2.2 漏洞二：密码明文展示 🔴 高危',
    '    2.3 漏洞三：CSRF 防护缺失 🟠 中危',
    '    2.4 漏洞四：敏感信息硬编码泄露 🟠 中危',
    '    2.5 漏洞五：生产环境 Debug 模式 🟡 低危',
    '3. 漏洞对比汇总',
    '4. 已确认安全的设计',
    '5. 修复建议总结',
    '6. 审计结论',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ===========================
# 1. 审计概述
# ===========================
add_colored_heading(doc, '1. 审计概述', RGBColor(0x1A, 0x1A, 0x2E), level=1)

doc.add_paragraph(
    '本次审计针对一个基于 Python Flask 框架开发的简易用户信息管理平台进行安全评估。'
    '该平台提供用户登录认证和个人信息展示功能。审计共发现 5 项安全漏洞，按严重程度分级如下：'
)

# 漏洞数量统计表
table = doc.add_table(rows=2, cols=2, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = table.rows[0].cells
cells[0].text = '严重程度'
cells[1].text = '数量'
set_cell_shading(cells[0], '1A1A2E')
set_cell_shading(cells[1], '1A1A2E')
for cell in cells:
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

cells = table.rows[1].cells
cells[0].text = '🔴 高危 / 🟠 中危 / 🟡 低危'
cells[1].text = '5 项（2 高危 + 2 中危 + 1 低危）'
for cell in cells:
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ===========================
# 2. 漏洞详情
# ===========================
add_colored_heading(doc, '2. 漏洞详情', RGBColor(0x1A, 0x1A, 0x2E), level=1)

# ---------- 漏洞一 ----------
doc.add_heading('2.1 漏洞一：密码明文存储', level=2)
add_risk_badge(doc, '🔴 高危 | CVSS 7.5 | CWE-256')

table = doc.add_table(rows=4, cols=2, style='Table Grid')
data = [
    ('漏洞编号', 'SEC-2026-001'),
    ('发现文件', 'app.py — USERS 字典'),
    ('漏洞类型', '敏感数据暴露（密码明文存储）'),
    ('违反标准', 'OWASP Top 10 A04、等保 2.0 三级身份鉴别'),
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    set_cell_shading(table.rows[i].cells[0], 'F0F0F0')
    table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

doc.add_paragraph(
    '用户密码以明文形式存储在 USERS 字典中，未做任何哈希处理。'
    '任何人获得源代码或数据库访问权限即可直接获取所有用户的明文密码。'
    '登录验证也使用 == 直接比对字符串，存在时序攻击风险。'
)

add_code_snippet(doc, '''USERS = {
    "admin": {
        "password": "admin123",  # ← 明文密码
    },
    "alice": {
        "password": "alice2025",  # ← 明文密码
    }
}

# 登录比对
if USERS[username]["password"] == password:  # ← 明文 == 对比''', label="原版漏洞代码")

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('风险分析：')
run.bold = True

risks = [
    '源码泄露 → 所有密码直接暴露',
    '用户常跨平台复用密码 → 可被撞库攻击',
    '违反密码存储安全规范，面临合规风险',
]
for r in risks:
    p = doc.add_paragraph(r, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('修复方案：')
run.bold = True

doc.add_paragraph(
    '使用 werkzeug.security.generate_password_hash() 进行 PBKDF2 哈希存储，'
    '登录时使用 check_password_hash() 进行安全比对。'
)

add_code_snippet(doc, '''from werkzeug.security import generate_password_hash, check_password_hash

USERS = {
    "admin": {
        "password": generate_password_hash("admin123"),  # ← 哈希存储
    },
    "alice": {
        "password": generate_password_hash("alice2025"),  # ← 哈希存储
    }
}

# 登录使用安全比对
if user and check_password_hash(user["password"], password):  # ← 哈希比对''', label="修复后代码")

doc.add_page_break()

# ---------- 漏洞二 ----------
doc.add_heading('2.2 漏洞二：密码明文展示', level=2)
add_risk_badge(doc, '🔴 高危 | CVSS 6.5 | CWE-200')

table = doc.add_table(rows=4, cols=2, style='Table Grid')
data = [
    ('漏洞编号', 'SEC-2026-002'),
    ('发现文件', 'app.py + templates/index.html'),
    ('漏洞类型', '敏感信息泄露（密码前端展示）'),
    ('违反标准', 'OWASP Top 10 A01、CWE-200'),
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    set_cell_shading(table.rows[i].cells[0], 'F0F0F0')
    table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
doc.add_paragraph(
    '用户登录成功后，密码被完整传递给前端模板，并直接渲染在 HTML 页面上。'
    '任何能接触到用户屏幕的人（肩窥攻击）均可获取密码。'
)

add_code_snippet(doc, '''# app.py - 将包含密码的完整用户信息传给模板
user_info = USERS[username]  # ← password 字段未过滤''', label="原版漏洞代码")

add_code_snippet(doc, '''<!-- index.html - 直接展示密码 -->
<li><span class="info-label">密码：</span>{{ user.password }}</li>''', label="原版漏洞代码（模板）")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('风险分析：')
run.bold = True

risks = [
    '肩窥攻击：他人从屏幕直接看到密码',
    '浏览器快照/页面缓存可能留存含密码的页面',
    '截屏分享时无意泄露凭证',
]
for r in risks:
    doc.add_paragraph(r, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('修复方案：')
run.bold = True
doc.add_paragraph('传递给模板时排除 password 字段，模板中删除密码展示行。')

add_code_snippet(doc, '''# app.py - 过滤密码字段
user_info = {k: v for k, v in user.items() if k != "password"}''', label="修复后代码")

add_code_snippet(doc, '''<!-- index.html - 移除密码展示 -->
<!-- password 行已删除 -->''', label="修复后代码（模板）")

doc.add_page_break()

# ---------- 漏洞三 ----------
doc.add_heading('2.3 漏洞三：CSRF 防护缺失', level=2)
add_risk_badge(doc, '🟠 中危 | CVSS 5.3 | CWE-352')

table = doc.add_table(rows=4, cols=2, style='Table Grid')
data = [
    ('漏洞编号', 'SEC-2026-003'),
    ('发现文件', 'app.py + templates/login.html'),
    ('漏洞类型', '跨站请求伪造（CSRF）'),
    ('违反标准', 'OWASP Top 10 A01、ASVS V12.1'),
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    set_cell_shading(table.rows[i].cells[0], 'F0F0F0')
    table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
doc.add_paragraph(
    '登录表单未包含 CSRF Token，也未对 POST 请求来源进行校验。'
    '攻击者可构造恶意页面诱导用户提交伪造的登录请求。'
)

add_code_snippet(doc, '''# app.py - GET 不生成 Token，POST 不校验来源
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "")
        # ... 直接处理，无 CSRF 校验''', label="原版漏洞代码")

add_code_snippet(doc, '''<!-- login.html - 表单无隐藏 CSRF 字段 -->
<form method="POST" action="/login">
    <input type="text" name="username">
    <input type="password" name="password">
    <button>登录</button>
</form>''', label="原版漏洞代码（模板）")

doc.add_paragraph()
risks = [
    '攻击者可构造跨站请求伪造用户登录',
    '结合会话固定攻击可能导致账号接管',
    '违背 OWASP Top 10 安全要求',
]
for r in risks:
    doc.add_paragraph(r, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('修复方案：')
run.bold = True
doc.add_paragraph('GET 请求时生成 secrets.token_hex(32) 存入 Session，表单添加隐藏 csrf_token 字段，POST 校验 Token。')

add_code_snippet(doc, '''import secrets

# GET 时生成 Token
session["csrf_token"] = secrets.token_hex(32)

# POST 时校验
csrf_token = request.form.get("csrf_token", "")
if csrf_token != session.get("csrf_token"):
    error = "表单已过期，请重新提交"''', label="修复后代码")

add_code_snippet(doc, '''<!-- login.html - 添加 CSRF 隐藏字段 -->
<input type="hidden" name="csrf_token" value="{{ csrf_token }}">''', label="修复后代码（模板）")

doc.add_page_break()

# ---------- 漏洞四 ----------
doc.add_heading('2.4 漏洞四：敏感信息硬编码泄露', level=2)
add_risk_badge(doc, '🟠 中危 | CVSS 4.0')

table = doc.add_table(rows=4, cols=2, style='Table Grid')
data = [
    ('漏洞编号', 'SEC-2026-004'),
    ('发现文件', 'templates/login.html'),
    ('漏洞类型', '敏感信息泄露（硬编码凭证）'),
    ('违反标准', 'CWE-200、等保 2.0 三级'),
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    set_cell_shading(table.rows[i].cells[0], 'F0F0F0')
    table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
doc.add_paragraph(
    '登录页面 HTML 源码的注释中直接写入了默认管理员账号和密码。'
    '任何查看页面源码的人都可以获取管理员凭证。'
)

add_code_snippet(doc, '''<!-- 调试信息 - 默认管理员账号 用户名: admin 密码: admin123 -->
{% extends "base.html" %}''', label="原版漏洞代码")

doc.add_paragraph()
risks = [
    '查看网页源码即可获取管理员密码',
    '浏览器开发者工具或 curl 直接可见',
    '搜索引擎可能索引到该敏感信息',
]
for r in risks:
    doc.add_paragraph(r, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('修复方案：')
run.bold = True
doc.add_paragraph('删除包含账号密码的 HTML 注释，默认凭证仅在文档中提及。')

# ---------- 漏洞五 ----------
doc.add_heading('2.5 漏洞五：生产环境 Debug 模式', level=2)
add_risk_badge(doc, '🟡 低危 | CVSS 3.5')

table = doc.add_table(rows=4, cols=2, style='Table Grid')
data = [
    ('漏洞编号', 'SEC-2026-005'),
    ('发现文件', 'app.py'),
    ('漏洞类型', '安全配置错误'),
    ('违反标准', 'OWASP Top 10 A05、CWE-489'),
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    set_cell_shading(table.rows[i].cells[0], 'F0F0F0')
    table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
doc.add_paragraph(
    'Flask 以 debug=True 模式运行，开启了 Werkzeug 调试器。'
    '该调试器在发生异常时显示完整堆栈跟踪，并在允许远程访问时可用于执行任意 Python 代码。'
)

add_code_snippet(doc, '''if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)  # ← debug=True''', label="原版漏洞代码")

doc.add_paragraph()
risks = [
    '调试器控制台可执行任意 Python 代码（需 PIN）',
    '堆栈跟踪泄露应用内部路径和变量',
    '应用异常时向用户暴露技术细节',
]
for r in risks:
    doc.add_paragraph(r, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('修复方案：')
run.bold = True
doc.add_paragraph('将 debug 设置为 False，或通过环境变量 FLASK_ENV 控制。')

add_code_snippet(doc, '''if __name__ == "__main__":
    app.run(debug=False, host="0.0.0.0", port=5000)  # ← debug=False''', label="修复后代码")

doc.add_page_break()

# ===========================
# 3. 漏洞对比汇总
# ===========================
add_colored_heading(doc, '3. 漏洞对比汇总', RGBColor(0x1A, 0x1A, 0x2E), level=1)

table = doc.add_table(rows=6, cols=5, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['漏洞名称', '原版', '修复版', '严重程度', 'CVSS']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '1A1A2E')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(10)

data = [
    ['密码明文存储', '❌ 明文 "admin123"', '✅ PBKDF2 哈希', '🔴 高危', '7.5'],
    ['密码明文展示', '❌ 页面显示密码', '✅ 已排除 password', '🔴 高危', '6.5'],
    ['CSRF 防护缺失', '❌ 无 Token 校验', '✅ 32 字节随机 Token', '🟠 中危', '5.3'],
    ['敏感信息硬编码', '❌ HTML 注释泄露', '✅ 已删除泄露注释', '🟠 中危', '4.0'],
    ['Debug 模式', '❌ debug=True', '✅ debug=False', '🟡 低危', '3.5'],
]
for i, row_data in enumerate(data):
    cells = table.rows[i + 1].cells
    for j, val in enumerate(row_data):
        cells[j].text = val
        for p in cells[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

# ===========================
# 4. 已确认安全的设计
# ===========================
add_colored_heading(doc, '4. 已确认安全的设计', RGBColor(0x1A, 0x1A, 0x2E), level=1)

table = doc.add_table(rows=5, cols=3, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['防护项', '状态', '说明']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, '1A1A2E')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

security_items = [
    ['XSS 防护', '✅ 有效', 'Jinja2 模板引擎默认开启 auto-escaping，HTML 特殊字符自动转义'],
    ['Session 管理', '✅ 有效', 'Flask 签名 Session，secret_key 已配置'],
    ['登出机制', '✅ 有效', '清除 username 和 csrf_token 两个 Session 键'],
    ['输入转义', '✅ 有效', '模板渲染时用户输入被正确转义，无法注入 HTML/JavaScript'],
]
for i, row_data in enumerate(security_items):
    cells = table.rows[i + 1].cells
    for j, val in enumerate(row_data):
        cells[j].text = val
        for p in cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

# ===========================
# 5. 修复建议总结
# ===========================
add_colored_heading(doc, '5. 修复建议总结', RGBColor(0x1A, 0x1A, 0x2E), level=1)

p = doc.add_paragraph()
run = p.add_run('紧急修复（高危）')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
urgent = [
    '对 USERS 字典中的密码使用 generate_password_hash() 进行哈希存储',
    '登录比对改用 check_password_hash() 安全函数',
    '用户信息传递到模板前过滤 password 字段',
    '模板中删除密码展示行',
]
for item in urgent:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('建议修复（中危）')
run.bold = True
run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
suggested = [
    '引入 secrets.token_hex() 生成 CSRF Token',
    '表单添加隐藏字段 {{ csrf_token }}',
    'POST 路由校验 Token 一致性',
    '删除 HTML 源码中的管理员凭证注释',
]
for item in suggested:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('最佳实践（低危）')
run.bold = True
run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
best = ['关闭 debug 模式或通过 FLASK_ENV 环境变量控制']
for item in best:
    doc.add_paragraph(item, style='List Bullet')

# ===========================
# 6. 审计结论
# ===========================
doc.add_page_break()
add_colored_heading(doc, '6. 审计结论', RGBColor(0x1A, 0x1A, 0x2E), level=1)

doc.add_paragraph(
    '原版项目存在 5 项安全漏洞，其中高危 2 项、中危 2 项、低危 1 项。'
    '主要安全问题集中在密码管理不当（明文存储 + 前端展示），'
    '以及缺少基本的 Web 安全防护机制（CSRF Token 缺失、Debug 模式开启）。'
)
doc.add_paragraph(
    '修复版已针对全部 5 项漏洞完成修复：'
    '密码采用 PBKDF2 哈希存储和比对，前端不再展示密码，'
    '增加了基于 Session 的 CSRF Token 校验机制，'
    '移除了 HTML 源码中的敏感信息，关闭了 Debug 模式。'
    '修复后代码达到基本安全水平，满足一般性 Web 应用的安全要求。'
)
doc.add_paragraph(
    '同时，Jinja2 模板引擎默认的 XSS 防护机制、Flask Session 管理、'
    '登出清理机制等在两个版本中均保持有效，无需额外修复。'
)

doc.add_paragraph()
doc.add_paragraph()

# 参考标准
p = doc.add_paragraph()
run = p.add_run('参考标准：')
run.bold = True

refs = [
    'OWASP Top 10 (2021)：A01（失效的访问控制）、A04（不安全的设计）、A05（安全配置错误）',
    'OWASP ASVS：V2.1（密码安全）、V3.1（会话管理）、V12.1（CSRF 防护）',
    '等保 2.0：三级安全要求 — 身份鉴别、访问控制、安全审计',
    'CWE：CWE-256（明文密码存储）、CWE-200（信息泄露）、CWE-352（CSRF）、CWE-489（Debug 后门）',
]
for ref in refs:
    doc.add_paragraph(ref, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 报告结束 —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(11)

# ===========================
# 保存
# ===========================
output_path = '/home/user/user-management/security-audit-report.docx'
doc.save(output_path)
print(f'✅ 报告已生成: {output_path}')
